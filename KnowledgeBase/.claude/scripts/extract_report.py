#!/usr/bin/env python3
"""
extract_report.py — Extract teks dari dokumen jadi teks mentah untuk
strukturkan jadi Markdown.

Sokongan: PDF (teks + OCR fallback), Word (.docx), Excel (.xlsx), imej.

Guna:
    python .claude/scripts/extract_report.py <path-fail> [--ocr]

Output: cetak teks mentah ke stdout (agent akan strukturkan jadi MD).

Install:
    pip install pdfplumber python-docx openpyxl pandas pdf2image pytesseract --break-system-packages
    # OCR (jika perlu): apt-get install -y tesseract-ocr poppler-utils
"""
import sys
import argparse
from pathlib import Path


def extract_pdf(path: Path, force_ocr: bool = False) -> str:
    out = []
    if not force_ocr:
        try:
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                for i, page in enumerate(pdf.pages, 1):
                    text = page.extract_text() or ""
                    out.append(f"\n--- Halaman {i} ---\n{text}")
                    for t_idx, table in enumerate(page.extract_tables() or [], 1):
                        out.append(f"\n[Jadual {i}.{t_idx}]")
                        for row in table:
                            cells = [str(c) if c is not None else "" for c in row]
                            out.append(" | ".join(cells))
            joined = "\n".join(out).strip()
            if joined and len(joined) > 30:
                return joined
        except Exception as e:
            print(f"[warn] pdfplumber gagal: {e}", file=sys.stderr)

    # OCR fallback untuk PDF imbasan
    try:
        from pdf2image import convert_from_path
        import pytesseract
        images = convert_from_path(str(path))
        ocr = []
        for i, img in enumerate(images, 1):
            ocr.append(f"\n--- Halaman {i} (OCR) ---\n{pytesseract.image_to_string(img)}")
        return "\n".join(ocr).strip()
    except Exception as e:
        return f"[ralat] Tidak dapat extract PDF (teks & OCR gagal): {e}"


def extract_docx(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    out = []
    for para in doc.paragraphs:
        if para.text.strip():
            out.append(para.text)
    for t_idx, table in enumerate(doc.tables, 1):
        out.append(f"\n[Jadual {t_idx}]")
        for row in table.rows:
            out.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(out).strip()


def extract_xlsx(path: Path) -> str:
    import pandas as pd
    out = []
    xls = pd.ExcelFile(path)
    for sheet in xls.sheet_names:
        df = xls.parse(sheet)
        out.append(f"\n--- Sheet: {sheet} ---")
        out.append(df.to_markdown(index=False))
    return "\n".join(out).strip()


def extract_image(path: Path) -> str:
    import pytesseract
    from PIL import Image
    return pytesseract.image_to_string(Image.open(path)).strip()


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("path")
    ap.add_argument("--ocr", action="store_true", help="Paksa OCR untuk PDF")
    args = ap.parse_args()

    path = Path(args.path)
    if not path.exists():
        print(f"[ralat] Fail tidak wujud: {path}", file=sys.stderr)
        sys.exit(1)

    ext = path.suffix.lower()
    if ext == ".pdf":
        print(extract_pdf(path, force_ocr=args.ocr))
    elif ext == ".docx":
        print(extract_docx(path))
    elif ext in (".xlsx", ".xlsm"):
        print(extract_xlsx(path))
    elif ext in (".png", ".jpg", ".jpeg", ".tiff", ".bmp"):
        print(extract_image(path))
    else:
        print(f"[ralat] Format tidak disokong: {ext}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
